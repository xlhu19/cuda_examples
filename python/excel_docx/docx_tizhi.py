# coding=utf-8
import docx
import random
import xlrd
from xlrd import xldate_as_datetime
from datetime import datetime

def read_names():
    data = xlrd.open_workbook("学生姓名2.2.xlsx")
    table = data.sheets()[0]
    ret = []
    for i in range(table.nrows):
        ret.append((table.col_values(0)[i], table.col_values(1)[i], table.col_values(2)[i], table.col_values(3)[i]))
    return ret

def update_select_row(_row):
    for i in [2,3,4,5,8,9,10,11]:
        if _row.cells[i].text != '':
            _row.cells[i].paragraphs[0].runs[0].text = ''
    _row.cells[random.randint(2,4)].text = '√'
    _row.cells[random.randint(8,10)].text = '√'



def generate_docs(data, ndoc):
    doc = docx.Document('tizhi_2.docx')
    # doc.paragraphs[2].runs[0].text = name

    dou = 4
    dou2 = 4
    dou3 = 4

    print(len(doc.tables))
    print(len(doc.tables[0].rows))
    doc.tables[0].rows[0].cells[1].paragraphs[0].runs[0].text = str(data[0])
    doc.tables[0].rows[0].cells[3].paragraphs[0].runs[0].text = str(data[1])
    doc.tables[0].rows[0].cells[5].paragraphs[0].runs[0].text = str(data[3])
    doc.tables[0].rows[1].cells[5].paragraphs[0].runs[0].text = str(xldate_as_datetime(data[2], 0).strftime("%Y-%m-%d"))

    grade_sum = 0

    bmi = random.uniform(12, 25)
    grade1 = 100
    grade2 = '正常'
    if bmi < 13.4:
        grade1 = 80
        grade2 = '偏瘦'
    elif bmi >= 18.2 and bmi <= 20.3:
        grade1 = 80
        grade2 = '偏胖'
    elif bmi > 20.4:
        grade1 = 60
        grade2 = '肥胖'
    else:
        grade1 = 100
        grade2 = '正常'
    doc.tables[1].rows[2].cells[1+dou].paragraphs[0].runs[0].text = str(round(bmi,1))
    doc.tables[1].rows[2].cells[2+dou].paragraphs[0].runs[0].text = str(grade1)
    doc.tables[1].rows[2].cells[3+dou].paragraphs[0].runs[0].text = grade2
    grade_sum += grade1
    
    fhl = random.randint(700,1700)
    if fhl < 1200:
        grade1 = 70
        grade2 = '及格'
    elif fhl >= 1200 and fhl <= 1400:
        grade1 = 85
        grade2 = '良好'
    else:
        grade1 = 100
        grade2 = '优秀'
    doc.tables[1].rows[3].cells[1+dou].paragraphs[0].runs[0].text = str(fhl)
    doc.tables[1].rows[3].cells[2+dou].paragraphs[0].runs[0].text = str(grade1)
    doc.tables[1].rows[3].cells[3+dou].paragraphs[0].runs[0].text = grade2
    grade_sum += grade1

    j50 = random.uniform(10.2, 11.8)
    if j50 > 10.8:
        grade1 = 60
        grade2 = '及格'
    elif j50 >= 10.5 and j50 <= 10.8:
        grade1 = 80
        grade2 = '良好'
    else:
        grade1 = 100
        grade2 = '优秀'
    doc.tables[1].rows[4].cells[1+dou].paragraphs[0].runs[0].text = str(round(j50, 1))
    doc.tables[1].rows[4].cells[2+dou].paragraphs[0].runs[0].text = str(grade1)
    doc.tables[1].rows[4].cells[3+dou].paragraphs[0].runs[0].text = grade2
    grade_sum += grade1

    tqq = random.uniform(5.4, 16.1)
    if tqq < 10:
        grade1 = 60
        grade2 = '及格'
    elif tqq >= 10 and tqq <= 12:
        grade1 = 80
        grade2 = '良好'
    else:
        grade1 = 100
        grade2 = '优秀'
    doc.tables[1].rows[5].cells[1+dou].paragraphs[0].runs[0].text = str(round(tqq, 1))
    doc.tables[1].rows[5].cells[2+dou].paragraphs[0].runs[0].text = str(grade1)
    doc.tables[1].rows[5].cells[3+dou].paragraphs[0].runs[0].text = grade2
    grade_sum += grade1

    ts = random.randint(55,109)
    if ts < 80:
        grade1 = 60
        grade2 = '及格'
    elif ts >= 80 and ts <= 93:
        grade1 = 80
        grade2 = '良好'
    else:
        grade1 = 100
        grade2 = '优秀'
    doc.tables[1].rows[6].cells[1+dou].paragraphs[0].runs[0].text = str(ts)
    doc.tables[1].rows[6].cells[2+dou].paragraphs[0].runs[0].text = str(grade1)
    doc.tables[1].rows[6].cells[3+dou].paragraphs[0].runs[0].text = grade2
    doc.tables[1].rows[9].cells[1+dou3].paragraphs[0].runs[0].text = str(ts)
    grade_sum += grade1

    grade_sum = round(grade_sum/5, 1)
    if grade_sum < 60:
        grade2 = '及格'
    elif grade_sum >= 60 and grade_sum < 65:
        grade2 = '良好'
    else:
        grade2 = '优秀'
    doc.tables[1].rows[7].cells[1+dou2].paragraphs[0].runs[0].text = str(grade_sum)
    doc.tables[1].rows[10].cells[1+dou2].paragraphs[0].runs[0].text = str(grade_sum)
    doc.tables[1].rows[11].cells[1+dou2].paragraphs[0].runs[0].text = grade2
    doc.save("tizhi2.docx")
    for element in doc.element.body:
        ndoc.element.body.append(element)
    
data = read_names()
print(data[0][0])
print(data[0][1])
print(xldate_as_datetime(data[0][2], 0).strftime("%Y-%m-%d"))

docNew = docx.Document("tizhi3.docx")
# generate_docs(data[0], docNew)
for i in data:
    generate_docs(i, docNew)
docNew.save("tizhi3.docx")
