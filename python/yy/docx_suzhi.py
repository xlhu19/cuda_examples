# coding=utf-8
import docx
import random
import xlrd

def update_select_row(_row):
    for i in [2,3,4,5,8,9,10,11]:
        if _row.cells[i].text != '':
            _row.cells[i].paragraphs[0].runs[0].text = ''
    _row.cells[random.randint(2,4)].text = '√'
    _row.cells[random.randint(8,10)].text = '√'

def update_table2(_table):

    for i in range(1, 15):
        update_select_row(_table.rows[i])
    comments = [
    '你其实是一个比较腼腆的孩子，你的作业，学习成绩都比上学期有进步。其实你的接受能力不错，如果你课上更专心一些，积极思考，踊跃发言，作业的质量更高些，那么你的进步会更大。',
    '你是一个思想上积极要求上进的孩子，性格文静，班队活动能积极参加，劳动积极做事认真，课上能专心听讲，作业及时完成。如果你的胆子再大点，课上勤于思考，勇于发言，你的进步会更快。',
    '你是一个聪明活泼、诚实善良的好孩子，能言善道，乐于助人。你的作文成绩比上学期有了一定的进步，学习上也自觉多了。如果你上课更专心点，作业、考试时更细心点，那么你的进步会更显著。',
    '你是一个比较听话的孩子，你作业上的字如同你的外表一样清秀，学习上比较自觉，成绩优良。如果你上课积极开动脑筋，大胆举手发言，课后加强阅读，你的作文水平将会有所提高。',
    '你是一个思想上积极要求上进的孩子，性格文静，班队活动能积极参加。你的作文成绩比上学期有了一定的进步，学习上也自觉多了。如果你的胆子再大点，勤于思考，勇于发言，你的进步会更快。',
    '你是一个比较听话的孩子，你作业上的字如同你的外表一样清秀，学习上比较自觉，成绩优良。如果你上课积极开动脑筋，大胆举手发言，课后加强阅读，你的作文水平将会有所提高。',
    '你很有个性，爱好广泛，体育委员工作负责，有较强的集体荣誉感，学习认真，课上专心听讲，作业清楚整洁。如果你再加强点课外阅读，课上加强思考，勇敢发表自己的见解，那你的进步会更大。',
    '你是一个文静的孩子，有较强的集体荣誉感，乐意为大家做事。学习认真，课堂上注意听讲，作业认真完成，成绩优良。如果你课堂上能积极思考，大胆发言，那你的学习成绩会更加稳定。',
    ]
    _table.rows[15].cells[1].paragraphs[0].runs[0].text = comments[random.randint(0, len(comments)-1)]

def read_names():
    data = xlrd.open_workbook("学生姓名2011.3.2.xlsx")
    table = data.sheets()[0]
    return table.col_values(0)

def generate_docs(name, ndoc):
    doc = docx.Document('xx_3.docx')
    doc.paragraphs[2].runs[0].text = name

    print(len(doc.tables))
    print(len(doc.tables[0].rows))
    sum = 0
    ir = [
            random.randint(10,19), random.randint(15,19),
            random.randint(10,19), random.randint(15,19),
            random.randint(10,19), random.randint(15,19),
            random.randint(25,39), random.randint(25,39),
            random.randint(10,19), random.randint(10,19),
            random.randint(10,19), random.randint(10,19),
            random.randint(35,59), random.randint(39,59),
            random.randint(75,100), random.randint(75,100),
            random.randint(75,100), random.randint(75,100),
            random.randint(75,100), random.randint(75,100),
            random.randint(75,100), random.randint(75,100),
            random.randint(75,100), random.randint(75,100),
            random.randint(75,100), random.randint(75,100),
            random.randint(75,100), random.randint(75,100),
            random.randint(75,100), random.randint(75,100),
         ]
    print(ir[23])
    doc.tables[0].rows[1].cells[2].paragraphs[0].runs[0].text = str(ir[0])
    doc.tables[0].rows[1].cells[3].paragraphs[0].runs[0].text = str(ir[1])

    doc.tables[0].rows[2].cells[2].paragraphs[0].runs[0].text = str(ir[2])
    doc.tables[0].rows[2].cells[3].paragraphs[0].runs[0].text = str(ir[3])

    doc.tables[0].rows[3].cells[2].paragraphs[0].runs[0].text = str(ir[4])
    doc.tables[0].rows[3].cells[3].paragraphs[0].runs[0].text = str(ir[5])

    doc.tables[0].rows[4].cells[2].paragraphs[0].runs[0].text = str(ir[6])
    doc.tables[0].rows[4].cells[3].paragraphs[0].runs[0].text = str(ir[7])
    for j in range(0, 8):
        sum += ir[j]
    print(sum)
    print(doc.tables[0].rows[1].cells[4].paragraphs[0].runs[0].text)
    if sum > 170:
        doc.tables[0].rows[1].cells[4].paragraphs[0].runs[0].text = '优秀'
    else:
        doc.tables[0].rows[1].cells[4].paragraphs[0].runs[0].text = '良好'

    doc.tables[0].rows[5].cells[2].paragraphs[0].runs[0].text = str(ir[8])
    doc.tables[0].rows[5].cells[3].paragraphs[0].runs[0].text = str(ir[9])

    doc.tables[0].rows[6].cells[2].paragraphs[0].runs[0].text = str(ir[10])
    doc.tables[0].rows[6].cells[3].paragraphs[0].runs[0].text = str(ir[11])

    doc.tables[0].rows[7].cells[2].paragraphs[0].runs[0].text = str(ir[12])
    doc.tables[0].rows[7].cells[3].paragraphs[0].runs[0].text = str(ir[13])

    sum = 0
    for j in range(8, 14):
        sum += ir[j]
    print(sum)
    if sum > 170:
        doc.tables[0].rows[2].cells[4].paragraphs[0].runs[0].text = '优秀'
    else:
        doc.tables[0].rows[2].cells[4].paragraphs[0].runs[0].text = '良好'


    # 英语
    print(ir[23])
    doc.tables[0].rows[8].cells[2].paragraphs[0].runs[0].text = str(ir[22])
    doc.tables[0].rows[8].cells[3].paragraphs[0].runs[0].text = str(ir[23])
    if (ir[22] + ir[23]) > 170:
        doc.tables[0].rows[8].cells[4].paragraphs[0].runs[0].text = '优秀'
    else:
        doc.tables[0].rows[8].cells[4].paragraphs[0].runs[0].text = '良好'
    doc.tables[0].rows[9].cells[2].paragraphs[0].runs[0].text = str(ir[24])
    doc.tables[0].rows[9].cells[3].paragraphs[0].runs[0].text = str(ir[25])
    if (ir[24] + ir[25]) > 170:
        doc.tables[0].rows[9].cells[4].paragraphs[0].runs[0].text = '优秀'
    else:
        doc.tables[0].rows[9].cells[4].paragraphs[0].runs[0].text = '良好'
    doc.tables[0].rows[14].cells[2].paragraphs[0].runs[0].text = str(ir[26])
    doc.tables[0].rows[14].cells[3].paragraphs[0].runs[0].text = str(ir[27])
    if (ir[26] + ir[27]) > 170:
        doc.tables[0].rows[14].cells[4].paragraphs[0].runs[0].text = '优秀'
    else:
        doc.tables[0].rows[14].cells[4].paragraphs[0].runs[0].text = '良好'
    doc.tables[0].rows[15].cells[2].paragraphs[0].runs[0].text = str(ir[28])
    doc.tables[0].rows[15].cells[3].paragraphs[0].runs[0].text = str(ir[29])
    if (ir[28] + ir[29]) > 170:
        doc.tables[0].rows[15].cells[4].paragraphs[0].runs[0].text = '优秀'
    else:
        doc.tables[0].rows[15].cells[4].paragraphs[0].runs[0].text = '良好'


    doc.tables[0].rows[10].cells[2].paragraphs[0].runs[0].text = str(ir[14])
    doc.tables[0].rows[10].cells[3].paragraphs[0].runs[0].text = str(ir[15])
    sum = 0
    for j in range(14, 16):
        sum += ir[j]
    print(sum)
    if sum > 170:
        doc.tables[0].rows[10].cells[4].paragraphs[0].runs[0].text = '优秀'
    else:
        doc.tables[0].rows[10].cells[4].paragraphs[0].runs[0].text = '良好'


    doc.tables[0].rows[11].cells[2].paragraphs[0].runs[0].text = str(ir[16])
    doc.tables[0].rows[11].cells[3].paragraphs[0].runs[0].text = str(ir[17])
    sum = 0
    for j in range(16, 18):
        sum += ir[j]
    print(sum)
    if sum > 170:
        doc.tables[0].rows[11].cells[4].paragraphs[0].runs[0].text = '优秀'
    else:
        doc.tables[0].rows[11].cells[4].paragraphs[0].runs[0].text = '良好'


    doc.tables[0].rows[12].cells[2].paragraphs[0].runs[0].text = str(ir[18])
    doc.tables[0].rows[12].cells[3].paragraphs[0].runs[0].text = str(ir[19])
    sum = 0
    for j in range(18, 20):
        sum += ir[j]
    print(sum)
    if sum > 170:
        doc.tables[0].rows[12].cells[4].paragraphs[0].runs[0].text = '优秀'
    else:
        doc.tables[0].rows[12].cells[4].paragraphs[0].runs[0].text = '良好'

    doc.tables[0].rows[13].cells[2].paragraphs[0].runs[0].text = str(ir[20])
    doc.tables[0].rows[13].cells[3].paragraphs[0].runs[0].text = str(ir[21])
    sum = 0
    for j in range(20, 22):
        sum += ir[j]
    print(sum)
    if sum > 170:
        doc.tables[0].rows[13].cells[4].paragraphs[0].runs[0].text = '优秀'
    else:
        doc.tables[0].rows[13].cells[4].paragraphs[0].runs[0].text = '良好'


    special = ('唱歌', '跳舞', '书法', '跳高', '美术', '音乐')
    doc.tables[0].rows[1].cells[5].paragraphs[0].runs[0].text = special[random.randint(0,5)]

    update_table2(doc.tables[1])


    doc.tables[0].rows[1].cells[5].paragraphs[0].runs[0].text = special[random.randint(0,5)]

    doc.save("_xx.docx")
    for element in doc.element.body:
        ndoc.element.body.append(element)
    
names = read_names()

docNew = docx.Document("xxx.docx")
# generate_docs(names[0], docNew)
for i in names:
    generate_docs(i, docNew)

docNew.save("xxx.docx")
