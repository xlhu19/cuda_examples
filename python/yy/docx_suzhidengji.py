# coding=utf-8
import docx
import random
import xlrd
from xlrd import xldate_as_datetime
from datetime import datetime

def read_names():
    data = xlrd.open_workbook("学生姓名2011.3.2.xlsx")
    table = data.sheets()[0]
    ret = []
    for i in range(table.nrows):
        ret.append((
            table.col_values(0)[i],
            table.col_values(1)[i],
            table.col_values(2)[i],
            table.col_values(3)[i],
            table.col_values(4)[i],
            table.col_values(5)[i],
            table.col_values(6)[i]
        ))
    return ret

def generate_docs(data, ndoc):
    doc = docx.Document('suzhidengji.docx')
    # print(doc.paragraphs[3].runs[0].text)
    print(doc.tables[0].rows[1].cells[16].paragraphs[0].runs[0].text)

    print(len(doc.paragraphs))
    print(len(doc.tables))
    print(len(doc.tables[0].rows))


    doc.tables[0].rows[1].cells[16].paragraphs[0].runs[0].text = str(int(data[3]))
    doc.tables[0].rows[2].cells[4].paragraphs[0].runs[0].text = data[0]
    doc.tables[0].rows[2].cells[10].paragraphs[0].runs[0].text = data[1]
    doc.tables[0].rows[3].cells[2].paragraphs[0].runs[0].text = data[4]
    doc.tables[0].rows[5].cells[0].paragraphs[0].runs[0].text = data[5]
    doc.tables[0].rows[5].cells[16].paragraphs[0].runs[0].text = str(int(data[6]))

    doc.save("suzhidengji2.docx")
    for element in doc.element.body:
        ndoc.element.body.append(element)
 
data = read_names()
print(xldate_as_datetime(data[0][2], 0).strftime("%Y-%m-%d"))

docNew = docx.Document("suzhidengji3.docx")
# generate_docs(data[0], docNew)
for i in data:
    generate_docs(i, docNew)

docNew.save("suzhidengji3.docx")
